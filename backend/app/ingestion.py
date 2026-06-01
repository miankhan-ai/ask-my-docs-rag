"""
Document ingestion pipeline: parse -> chunk -> embed -> store.

Public functions
----------------
parse_file        : Parse PDF / DOCX / Markdown / plain text into page dicts.
chunk_pages       : Split parsed pages into overlapping chunks with char offsets.
embed_chunks      : Embed chunk texts via the HuggingFace Inference API (batched).
ingest_document   : End-to-end: persist a Document + its embedded Chunks.
"""

import tempfile
from pathlib import Path
from typing import Any

import pypdf
import docx as python_docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.embeddings import embed_texts
from app.models import Document, DocumentStatus, Chunk


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #
def parse_file(path: Path) -> list[dict[str, Any]]:
    """Parse an uploaded file into a list of ``{text, page_number}`` dicts."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix in (".md", ".txt"):
        return _parse_text(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> list[dict[str, Any]]:
    pages = []
    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"text": text, "page_number": i + 1})
    return pages


def _parse_docx(path: Path) -> list[dict[str, Any]]:
    doc = python_docx.Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [{"text": text, "page_number": None}]


def _parse_text(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8")
    return [{"text": text, "page_number": None}]


# --------------------------------------------------------------------------- #
# Chunking
# --------------------------------------------------------------------------- #
def chunk_pages(
    pages: list[dict[str, Any]],
    chunk_size: int = 512,
    chunk_overlap: int = 64,
) -> list[dict[str, Any]]:
    """Split parsed pages into overlapping chunks with char-offset metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        add_start_index=True,
    )
    chunks: list[dict[str, Any]] = []
    chunk_index = 0
    for page in pages:
        text = page["text"]
        docs = splitter.create_documents([text])
        for doc in docs:
            start = doc.metadata.get("start_index", 0)
            end = start + len(doc.page_content)
            chunks.append(
                {
                    "text": doc.page_content,
                    "page_number": page["page_number"],
                    "char_start": start,
                    "char_end": end,
                    "chunk_index": chunk_index,
                }
            )
            chunk_index += 1
    return chunks


# --------------------------------------------------------------------------- #
# Embedding
# --------------------------------------------------------------------------- #
async def embed_chunks(
    chunks: list[dict[str, Any]],
    batch_size: int | None = None,
) -> list[list[float]]:
    """Embed chunk texts in batches; returns a list of 1024-dim vectors.

    Uses the configured embedding backend (HF Inference API or local
    sentence-transformers) via :func:`app.embeddings.embed_texts`.
    """
    if batch_size is None:
        batch_size = settings.embedding_batch_size
    texts = [c["text"] for c in chunks]
    embeddings: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        batch_embeddings = await embed_texts(batch)
        embeddings.extend(batch_embeddings)
    return embeddings


# --------------------------------------------------------------------------- #
# End-to-end ingestion
# --------------------------------------------------------------------------- #
async def ingest_document(
    db: AsyncSession,
    filename: str,
    filetype: str,
    file_bytes: bytes,
) -> tuple[int, int]:
    """Parse, chunk, embed and store a document. Returns ``(document_id, chunk_count)``."""
    doc = Document(filename=filename, filetype=filetype, status=DocumentStatus.processing)
    db.add(doc)
    await db.flush()

    try:
        with tempfile.NamedTemporaryFile(suffix=filetype, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = Path(tmp.name)

        pages = parse_file(tmp_path)
        tmp_path.unlink(missing_ok=True)

        chunks_meta = chunk_pages(
            pages,
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )
        embeddings = await embed_chunks(chunks_meta)

        chunk_objects = [
            Chunk(
                document_id=doc.id,
                text=meta["text"],
                page_number=meta["page_number"],
                char_start=meta["char_start"],
                char_end=meta["char_end"],
                chunk_index=meta["chunk_index"],
                embedding=emb,
            )
            for meta, emb in zip(chunks_meta, embeddings)
        ]
        db.add_all(chunk_objects)
        doc.status = DocumentStatus.ready
        await db.commit()
        return doc.id, len(chunk_objects)

    except Exception:
        doc.status = DocumentStatus.failed
        await db.commit()
        raise
